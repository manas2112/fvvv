import telebot
from telebot import types
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
import os
from datetime import datetime
import tempfile
from PIL import Image
import io
import shutil

# Bot tokenini o'rnating
TOKEN = '8393199376:AAEpm8oSALlRRKTBX9KWOojB6W56GUem1pE'
bot = telebot.TeleBot(TOKEN)

# Foydalanuvchilar holatini saqlash
user_data = {}
user_states = {}

# Holatlar
STATE_DATE = 1
STATE_NAME = 2
STATE_PHOTO = 3

@bot.message_handler(commands=['start'])
def start(message):
    chat_id = message.chat.id
    user_states[chat_id] = STATE_DATE
    
    # Yangi foydalanuvchi uchun ma'lumotlar strukturasini yaratish
    user_data[chat_id] = {
        'records': [],  # Barcha saqlangan yozuvlar
        'current_record': {  # Hozirgi yozuv
            'date': '',
            'name': '',
            'photos': []
        }
    }
    
    markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
    btn1 = types.KeyboardButton("✅ Yakunlash")
    btn2 = types.KeyboardButton("📊 Hisobot")
    markup.add(btn1, btn2)
    
    bot.send_message(
        chat_id, 
        "Assalomu alaykum! FVV xodimi uchun bot.\n\n"
        "Iltimos, tekshiruv sanasini kiriting (KK.OO.YYYY formatida):",
        reply_markup=markup
    )

@bot.message_handler(func=lambda message: message.text == "✅ Yakunlash")
def finish_and_create_report(message):
    chat_id = message.chat.id
    
    # Hozirgi yozuvni tekshirish va saqlash
    current = user_data[chat_id]['current_record']
    
    # Agar hozirgi yozuvda ma'lumotlar bo'lsa, uni saqlash
    if current.get('date') and current.get('name') and current.get('photos'):
        user_data[chat_id]['records'].append(current.copy())
        # Hozirgi yozuvni tozalash
        user_data[chat_id]['current_record'] = {
            'date': '',
            'name': '',
            'photos': []
        }
    
    # Jami yozuvlarni tekshirish
    if not user_data[chat_id]['records']:
        bot.send_message(chat_id, "❌ Hali hech qanday ma'lumot kiritilmadi!\n\nSana va ism familiya kiriting, rasm yuboring.")
        user_states[chat_id] = STATE_DATE
        return
    
    try:
        # Word fayl yaratish
        filename = create_word_report(user_data[chat_id]['records'], chat_id)
        
        # Faylni yuborish
        with open(filename, 'rb') as file:
            bot.send_document(chat_id, file, caption="✅ Hisobot tayyor!")
        
        # Faylni o'chirish
        os.remove(filename)
        
        # Ma'lumotlarni tozalash
        user_data[chat_id]['records'] = []
        user_states[chat_id] = STATE_DATE
        
        bot.send_message(
            chat_id, 
            "✅ Hisobot yuborildi!\n\n"
            "Yangi ma'lumot kiritish uchun sanani kiriting:"
        )
        
    except Exception as e:
        bot.send_message(chat_id, f"❌ Xatolik yuz berdi: {str(e)}")

@bot.message_handler(func=lambda message: message.text == "📊 Hisobot")
def show_stats(message):
    chat_id = message.chat.id
    
    if chat_id not in user_data:
        count = 0
    else:
        count = len(user_data[chat_id]['records'])
        # Hozirgi yozuvni ham hisobga olish
        if user_data[chat_id]['current_record'].get('name'):
            count += 1
    
    bot.send_message(
        chat_id, 
        f"📊 Statistika:\n"
        f"Jami kiritilgan uylar: {count}\n"
        f"Oxirgi tekshiruv: {datetime.now().strftime('%d.%m.%Y')}"
    )

@bot.message_handler(func=lambda message: True)
def handle_messages(message):
    chat_id = message.chat.id
    text = message.text
    
    # Yakunlash va Hisobot tugmalari bosilganda alohida ishlov beriladi
    if text in ["✅ Yakunlash", "📊 Hisobot"]:
        return
    
    # Foydalanuvchi ma'lumotlarini tekshirish
    if chat_id not in user_data:
        user_data[chat_id] = {
            'records': [],
            'current_record': {
                'date': '',
                'name': '',
                'photos': []
            }
        }
        user_states[chat_id] = STATE_DATE
    
    state = user_states.get(chat_id, STATE_DATE)
    
    if state == STATE_DATE:
        # Sanani tekshirish
        try:
            # Sanani tekshirish
            date_obj = datetime.strptime(text, '%d.%m.%Y')
            current_date = date_obj.strftime('%d.%m.%Y')
            
            # Yangi yozuv boshlash
            user_data[chat_id]['current_record'] = {
                'date': current_date,
                'name': '',
                'photos': []
            }
            
            user_states[chat_id] = STATE_NAME
            bot.send_message(chat_id, "✅ Sana saqlandi!\nEndi uy egasining ism-familiyasini kiriting:")
            
        except ValueError:
            bot.send_message(
                chat_id, 
                "❌ Noto'g'ri format! Iltimos, sanani KK.OO.YYYY formatida kiriting.\n"
                "Misol: 25.01.2026"
            )
    
    elif state == STATE_NAME:
        # Ism familiyani saqlash
        if len(text) < 3:
            bot.send_message(chat_id, "❌ Ism familiya juda qisqa. Qayta kiriting:")
            return
        
        user_data[chat_id]['current_record']['name'] = text
        user_states[chat_id] = STATE_PHOTO
        
        bot.send_message(
            chat_id, 
            f"✅ Ism familiya saqlandi: {text}\n\n"
            f"📸 Endi katol va gaz plitasining rasmini yuboring.\n"
            f"(Bir nechta rasm yuborishingiz mumkin)\n\n"
            f"Rasm yuborganingizdan so'ng, yangi uy uchun ma'lumot kiritishni davom ettirishingiz yoki "
            f"'✅ Yakunlash' tugmasini bosib hisobotni olishingiz mumkin."
        )
    
    else:  # STATE_PHOTO holatida
        bot.send_message(
            chat_id, 
            "❌ Iltimos, avval rasm yuboring!\n"
            "Yoki yangi uy kiritish uchun sanani kiriting."
        )

@bot.message_handler(content_types=['photo'])
def handle_photo(message):
    chat_id = message.chat.id
    
    # Foydalanuvchi ma'lumotlarini tekshirish
    if chat_id not in user_data:
        user_data[chat_id] = {
            'records': [],
            'current_record': {
                'date': '',
                'name': '',
                'photos': []
            }
        }
        user_states[chat_id] = STATE_DATE
        bot.send_message(chat_id, "❌ Avval sanani kiriting! /start ni bosing.")
        return
    
    # Hozirgi yozuvni tekshirish
    if not user_data[chat_id]['current_record'].get('date'):
        bot.send_message(chat_id, "❌ Avval sanani kiriting!")
        user_states[chat_id] = STATE_DATE
        return
    
    if not user_data[chat_id]['current_record'].get('name'):
        bot.send_message(chat_id, "❌ Avval ism familiyani kiriting!")
        user_states[chat_id] = STATE_NAME
        return
    
    try:
        # Rasmni olish
        file_id = message.photo[-1].file_id
        file_info = bot.get_file(file_id)
        downloaded_file = bot.download_file(file_info.file_path)
        
        # Rasmni saqlash
        if 'photos' not in user_data[chat_id]['current_record']:
            user_data[chat_id]['current_record']['photos'] = []
        
        user_data[chat_id]['current_record']['photos'].append(downloaded_file)
        
        # Rasm sonini ko'rsatish
        photo_count = len(user_data[chat_id]['current_record']['photos'])
        
        markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
        btn1 = types.KeyboardButton("✅ Uyni yakunlash")
        btn2 = types.KeyboardButton("✅ Yakunlash")
        markup.add(btn1, btn2)
        bot.send_message(
            chat_id, 
            f"✅ Rasm qabul qilindi! ({photo_count} ta rasm)\n\n"
            f"📋 Joriy uy ma'lumotlari:\n"
            f"📅 Sana: {user_data[chat_id]['current_record']['date']}\n"
            f"👤 Egasi: {user_data[chat_id]['current_record']['name']}\n"
            f"📸 Rasmlar: {photo_count} ta\n\n"
            f"Yana rasm yuborishingiz mumkin yoki:\n"
            f"• '✅ Uyni yakunlash' - bu uy uchun kiritishni tugatish\n"
            f"• '✅ Yakunlash' - umumiy hisobotni olish",
            reply_markup=markup
        )
        
    except Exception as e:
        bot.send_message(chat_id, f"❌ Rasm saqlashda xatolik: {str(e)}")

@bot.message_handler(func=lambda message: message.text == "✅ Uyni yakunlash")
def finish_house(message):
    chat_id = message.chat.id
    
    if chat_id not in user_data:
        bot.send_message(chat_id, "❌ Ma'lumot topilmadi! /start ni bosing.")
        return
    
    current = user_data[chat_id]['current_record']
    
    # Tekshirish
    if not current.get('date'):
        bot.send_message(chat_id, "❌ Sana kiritilmagan!")
        user_states[chat_id] = STATE_DATE
        return
    
    if not current.get('name'):
        bot.send_message(chat_id, "❌ Ism familiya kiritilmagan!")
        user_states[chat_id] = STATE_NAME
        return
    
    if not current.get('photos'):
        bot.send_message(chat_id, "❌ Hech qanday rasm kiritilmagan!")
        return
    
    # Ma'lumotni saqlash
    user_data[chat_id]['records'].append(current.copy())
    
    # Hozirgi yozuvni tozalash
    user_data[chat_id]['current_record'] = {
        'date': '',
        'name': '',
        'photos': []
    }
    
    user_states[chat_id] = STATE_DATE
    
    # Asosiy menyuni qaytarish
    markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
    btn1 = types.KeyboardButton("✅ Yakunlash")
    btn2 = types.KeyboardButton("📊 Hisobot")
    markup.add(btn1, btn2)
    
    bot.send_message(
        chat_id, 
        f"✅ Uy ma'lumotlari saqlandi!\n"
        f"Jami kiritilgan uylar: {len(user_data[chat_id]['records'])}\n\n"
        f"Yangi uy uchun sanani kiriting:",
        reply_markup=markup
    )

def create_word_report(records, chat_id):
    """Word hisobot yaratish - barcha rasmlar bir sahifada"""
    
    # Yangi Word hujjat yaratish
    doc = Document()
    
    # Sahifa sozlamalari (chekka joylarni kichikroq qilish)
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    
    # Sarlavha
    title = doc.add_heading('FVV TEKSHIRUV HISOBOTI', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Umumiy ma'lumot
    doc.add_paragraph(f"Hisobot sanasi: {datetime.now().strftime('%d.%m.%Y')}")
    doc.add_paragraph(f"Tekshirilgan uylar soni: {len(records)}")
    doc.add_paragraph("=" * 50)
    
    # Har bir uy uchun alohida bo'lim
    for idx, record in enumerate(records, 1):
        # Yangi sahifadan boshlash (1-sahifadan tashqari)
        if idx > 1:
            doc.add_page_break()
        
        # Uy raqami sarlavhasi
        doc.add_heading(f"UY №{idx}", level=1)
        
        # Ma'lumotlar
        doc.add_paragraph(f"📅 Tekshiruv sanasi: {record['date']}")
        doc.add_paragraph(f"👤 Uy egasi: {record['name']}")
        doc.add_paragraph(f"📸 Rasmlar soni: {len(record.get('photos', []))}")
        
        doc.add_paragraph("-" * 40)
        
        # Rasmlarni qo'shish - barchasi bir sahifada
        if record.get('photos') and len(record['photos']) > 0:
            doc.add_heading('Tekshiruv rasmlari:', level=2)
            
            # Vaqtinchalik papka yaratish
            temp_dir = tempfile.mkdtemp()
            
            # Rasm soniga qarab jadval o'lchamini aniqlash
            photos_count = len(record['photos'])
            
            if photos_count <= 2:
                # 1 qator, 2 ustun
                table = doc.add_table(rows=1, cols=2)
                rows, cols = 1, 2
            elif photos_count <= 4:
                # 2 qator, 2 ustun
                table = doc.add_table(rows=2, cols=2)
                rows, cols = 2, 2
            elif photos_count <= 6:
                # 3 qator, 2 ustun
                table = doc.add_table(rows=3, cols=2)
                rows, cols = 3, 2
            elif photos_count <= 9:
                # 3 qator, 3 ustun
                table = doc.add_table(rows=3, cols=3)
                rows, cols = 3, 3
            else:
                # 4 qator, 3 ustun (12 ta rasm)
                table = doc.add_table(rows=4, cols=3)
                rows, cols = 4, 3
            
            # Jadval uslubi
            table.style = 'Table Grid'
            
            # Ustun kengliklarini sozlash
            for col in table.columns:
                for cell in col.cells:
                    cell.width = Inches(2.5)
                    # Hujayra chetlarini sozlash
                    cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Rasmlarni jadvalga joylashtirish
            photo_index = 0
            for row in range(rows):
                for col in range(cols):
                    if photo_index < photos_count:
                        try:
                            # Rasmni vaqtinchalik faylga saqlash
                            img_path = os.path.join(temp_dir, f"photo_{photo_index}.jpg")
                            with open(img_path, 'wb') as f:
                                f.write(record['photos'][photo_index])
                            
                            # Rasmni o'lchamini optimallashtirish
                            img = Image.open(img_path)
                            
                            # Jadval katakchasiga rasm qo'shish
                            cell = table.cell(row, col)
                            
                            # Katakdagi matnni tozalash
                            cell.text = ''
                            
                            # Rasm qo'shish
                            paragraph = cell.add_paragraph()
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = paragraph.add_run()
                            
                            # Rasm hajmini katakka moslash
                            if photos_count <= 4:
                                run.add_picture(img_path, width=Inches(2.2), height=Inches(1.8))
                            else:
                                run.add_picture(img_path, width=Inches(1.8), height=Inches(1.5))
                            
                            # Rasm tagiga raqam qo'shish
                            caption = cell.add_paragraph()
                            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            caption.add_run(f"Rasm {photo_index + 1}").bold = True
                            caption.paragraph_format.space_before = Pt(2)
                            
                            photo_index += 1
                            
                        except Exception as e:
                            print(f"Rasm qo'shishda xatolik: {e}")
                            photo_index += 1
                    else:
                        # Bo'sh katakchalarni to'ldirish
                        cell = table.cell(row, col)
                        cell.text = ''
            
            # Vaqtinchalik fayllarni tozalash
            shutil.rmtree(temp_dir, ignore_errors=True)
            
            # Jadvaldan keyin bo'sh joy
            doc.add_paragraph()
        
        # Har bir uy oxirida chiziq
        doc.add_paragraph("=" * 50)
    
    # Faylni saqlash
    filename = f"fvv_hisobot_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)
    
    return filename
# Botni ishga tushirish
if __name__ == '__main__':
    print("Bot ishga tushdi...")
    print("Bot tokeni:", TOKEN)
    print("Kutubxonalar o'rnatilganligini tekshiring: python-docx, pillow")
    print("Kutubxonalarni o'rnatish: pip install pytelegrambotapi python-docx pillow")
    bot.infinity_polling()
